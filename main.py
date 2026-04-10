#!/usr/bin/env python3
"""PDF 批量转 Word 工具 — 双引擎：智能提取（PyMuPDF）/ 高保真（pdf2docx）"""

import os
import re
import platform
import subprocess
import threading
from pathlib import Path
import tkinter as tk
from tkinter import ttk, filedialog, messagebox

try:
    from tkinterdnd2 import DND_FILES, TkinterDnD
    HAS_DND = True
except ImportError:
    HAS_DND = False

# ── 配色 & 字体 ────────────────────────────────────────────────────────────
BG      = "#f0f2f5"
CARD    = "white"
BORDER  = "#dde1e7"
ACCENT  = "#4a90e2"
SUCCESS = "#28a745"
GRAY    = "#6c757d"
FONT    = "微软雅黑"


# ══════════════════════════════════════════════════════════════════════════════
#  智能提取引擎（PyMuPDF）
# ══════════════════════════════════════════════════════════════════════════════

LIBREOFFICE_PATHS = [
    "/Applications/LibreOffice.app/Contents/MacOS/soffice",   # macOS
    "/Applications/libreoffice.app/Contents/MacOS/soffice",
    "soffice",                                                  # Linux / PATH
    r"C:\Program Files\LibreOffice\program\soffice.exe",       # Windows
    r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
]


def _find_soffice() -> str:
    """返回可用的 soffice 可执行文件路径，找不到则抛出 RuntimeError。"""
    import shutil
    for p in LIBREOFFICE_PATHS:
        if os.path.isfile(p):
            return p
        if shutil.which(p):
            return p
    raise RuntimeError(
        "未找到 LibreOffice。\n"
        "请先安装：https://www.libreoffice.org/download/\n"
        "macOS 也可运行：brew install --cask libreoffice"
    )


def convert_libreoffice(pdf_path: str, docx_path: str) -> None:
    """调用 LibreOffice headless 模式将 PDF 转换为 DOCX。"""
    import subprocess, shutil, tempfile

    soffice = _find_soffice()
    dest_dir = str(Path(docx_path).parent)

    # LibreOffice 会在 outdir 下生成同名 .docx，先用临时目录接收
    with tempfile.TemporaryDirectory() as tmp:
        result = subprocess.run(
            [soffice, "--headless", "--norestore",
             "--convert-to", "docx",
             "--outdir", tmp,
             pdf_path],
            capture_output=True, text=True, timeout=120
        )
        if result.returncode != 0:
            raise RuntimeError(result.stderr.strip() or "LibreOffice 转换失败")

        # 找到生成的 docx 并移动到目标位置
        generated = list(Path(tmp).glob("*.docx"))
        if not generated:
            raise RuntimeError("LibreOffice 未生成输出文件")
        shutil.move(str(generated[0]), docx_path)


def convert_smart(pdf_path: str, docx_path: str) -> None:
    """
    用 PyMuPDF 做结构化提取：
    - 按字号/粗体识别标题层级
    - 内置表格检测（find_tables）
    - 图片以合理 DPI 嵌入
    - 输出干净的可编辑段落，接近 Adobe 的转换风格
    """
    import fitz
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import io

    pdf = fitz.open(pdf_path)
    word = Document()

    # ── Word 页面设置 ─────────────────────────────────────────────────────
    section = word.sections[0]
    section.page_width  = int(8.27 * 914400)   # A4
    section.page_height = int(11.69 * 914400)
    section.left_margin = section.right_margin = int(1.0 * 914400)
    section.top_margin  = section.bottom_margin = int(1.0 * 914400)

    # ── 段落默认样式 ──────────────────────────────────────────────────────
    normal = word.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)

    def _set_run_font(run, font_name: str, size_pt: float, bold: bool, italic: bool, color):
        run.font.name = font_name
        run.font.size = Pt(size_pt)
        run.bold      = bold
        run.italic    = italic
        if color:
            r, g, b = color
            if (r, g, b) != (0, 0, 0):
                run.font.color.rgb = RGBColor(r, g, b)

    def _add_paragraph(text: str, style: str = "Normal",
                       bold: bool = False, size: float = 11,
                       align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
        p = word.add_paragraph(style=style)
        p.alignment = align
        run = p.add_run(text)
        run.bold      = bold
        run.font.size = Pt(size)

    # ── 每页处理 ──────────────────────────────────────────────────────────
    for page_num, page in enumerate(pdf):
        page_rect = page.rect
        page_width = page_rect.width

        # 收集该页所有文字块的字号分布，用于相对判断标题
        spans_on_page = []
        raw = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE |
                                          fitz.TEXT_MEDIABOX_CLIP)
        for blk in raw["blocks"]:
            if blk["type"] != 0:
                continue
            for line in blk["lines"]:
                for sp in line["spans"]:
                    if sp["text"].strip():
                        spans_on_page.append(sp["size"])

        body_size = sorted(spans_on_page)[len(spans_on_page)//2] if spans_on_page else 11

        # ── 检测表格区域 ───────────────────────────────────────────────────
        finder = page.find_tables()
        table_rects = [t.bbox for t in finder.tables]

        def _in_table(rect) -> bool:
            rx0, ry0, rx1, ry1 = rect
            for tx0, ty0, tx1, ty1 in table_rects:
                if rx0 >= tx0 - 2 and ry0 >= ty0 - 2 and rx1 <= tx1 + 2 and ry1 <= ty1 + 2:
                    return True
            return False

        # ── 插入表格 ───────────────────────────────────────────────────────
        inserted_tables = set()
        for tbl_obj in finder.tables:
            tid = id(tbl_obj)
            if tid in inserted_tables:
                continue
            inserted_tables.add(tid)
            try:
                extracted = tbl_obj.extract()  # list[list[str|None]]
                if not extracted:
                    continue
                rows = len(extracted)
                cols = max(len(r) for r in extracted)
                if rows == 0 or cols == 0:
                    continue
                w_tbl = word.add_table(rows=rows, cols=cols)
                w_tbl.style = "Table Grid"
                for ri, row in enumerate(extracted):
                    for ci, cell_text in enumerate(row):
                        if ci < cols:
                            cell = w_tbl.cell(ri, ci)
                            cell.text = str(cell_text or "")
                            if ri == 0:
                                for run in cell.paragraphs[0].runs:
                                    run.bold = True
                word.add_paragraph()  # 表格后空行
            except Exception:
                pass

        # ── 处理文字块 ────────────────────────────────────────────────────
        prev_y1 = 0.0
        for blk in raw["blocks"]:
            if blk["type"] == 1:
                # 图片块
                try:
                    clip = fitz.Rect(blk["bbox"])
                    mat  = fitz.Matrix(2, 2)        # 2× 缩放提高清晰度
                    pix  = page.get_pixmap(matrix=mat, clip=clip, alpha=False)
                    img_bytes = pix.tobytes("png")
                    p = word.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    run.add_picture(io.BytesIO(img_bytes),
                                    width=Inches(min(6.0, (clip.width / page_width) * 6.27)))
                except Exception:
                    pass
                continue

            if blk["type"] != 0:
                continue

            blk_rect = blk["bbox"]  # (x0, y0, x1, y1)

            # 跳过落在表格内部的文字块
            if _in_table(blk_rect):
                continue

            # 检查段间空白，添加空行
            gap = blk_rect[1] - prev_y1
            if prev_y1 > 0 and gap > body_size * 1.5:
                word.add_paragraph()
            prev_y1 = blk_rect[3]

            # 合并块内所有行的文字
            block_text = ""
            block_spans = []
            for line in blk["lines"]:
                line_text = "".join(sp["text"] for sp in line["spans"])
                block_text += line_text + "\n"
                block_spans.extend(line["spans"])

            block_text = block_text.strip()
            if not block_text:
                continue

            if not block_spans:
                _add_paragraph(block_text)
                continue

            # 主要字号和样式（取第一个非空 span）
            primary = next((s for s in block_spans if s["text"].strip()), block_spans[0])
            size_pt = primary["size"]
            is_bold = bool(primary["flags"] & 2**4)

            # 居中判断（块横跨页面 >60% 且 x0 有缩进）
            bx0, bx1 = blk_rect[0], blk_rect[2]
            blk_w = bx1 - bx0
            is_center = (blk_w / page_width > 0.4) and (bx0 > page_width * 0.1)
            align = WD_ALIGN_PARAGRAPH.CENTER if is_center else WD_ALIGN_PARAGRAPH.LEFT

            # 标题判断
            ratio = size_pt / body_size if body_size > 0 else 1.0

            if ratio >= 1.5 or (ratio >= 1.2 and is_bold):
                style = "Heading 1"
            elif ratio >= 1.2 or (ratio >= 1.05 and is_bold):
                style = "Heading 2"
            elif is_bold and ratio >= 0.95:
                style = "Heading 3"
            else:
                style = "Normal"

            p = word.add_paragraph(style=style)
            p.alignment = align

            # 逐 span 插入，保留粗体/斜体
            for sp in block_spans:
                sp_text = sp["text"]
                if not sp_text:
                    continue
                run = p.add_run(sp_text)
                sp_bold   = bool(sp["flags"] & 2**4)
                sp_italic = bool(sp["flags"] & 2**1)
                color_int = sp.get("color", 0)
                color_rgb = ((color_int >> 16) & 0xFF,
                             (color_int >> 8)  & 0xFF,
                             color_int         & 0xFF)
                _set_run_font(run, "Calibri", sp["size"], sp_bold, sp_italic, color_rgb)

        # 分页符（最后一页不加）
        if page_num < len(pdf) - 1:
            word.add_page_break()

    pdf.close()
    word.save(docx_path)


# ══════════════════════════════════════════════════════════════════════════════
#  主 GUI
# ══════════════════════════════════════════════════════════════════════════════

class App:
    def __init__(self):
        self._check_deps()

        # tkinterdnd2 在 PyInstaller 打包后可能因 tkdnd 原生库缺失而失败
        # 此处做运行时降级处理
        dnd_ok = False
        if HAS_DND:
            try:
                self.root = TkinterDnD.Tk()
                dnd_ok = True
            except Exception:
                self.root = tk.Tk()
        else:
            self.root = tk.Tk()
        self._dnd_ok = dnd_ok
        self.root.title("PDF 批量转 Word 工具")
        self.root.configure(bg=BG)
        self.root.minsize(640, 600)
        self._center(760, 700)

        self.files: list = []
        self.out_dir  = tk.StringVar()
        self.engine   = tk.StringVar(value="smart")   # "smart" | "fidelity" | "libreoffice"
        self.running  = False

        self._setup_styles()
        self._build_ui()
        self.root.mainloop()

    # ── 依赖检查 ──────────────────────────────────────────────────────────────

    @staticmethod
    def _check_deps():
        missing = []
        try:
            import fitz  # noqa: F401
        except ImportError:
            missing.append("pymupdf")
        try:
            from docx import Document  # noqa: F401
        except ImportError:
            missing.append("python-docx")
        if missing:
            root = tk.Tk(); root.withdraw()
            messagebox.showerror("缺少依赖",
                f"缺少：{', '.join(missing)}\n\n请运行：pip install {' '.join(missing)}")
            root.destroy(); raise SystemExit(1)

    def _center(self, w, h):
        self.root.geometry(f"{w}x{h}")
        self.root.update_idletasks()
        x = (self.root.winfo_screenwidth()  - w) // 2
        y = (self.root.winfo_screenheight() - h) // 2
        self.root.geometry(f"{w}x{h}+{x}+{y}")

    def _setup_styles(self):
        s = ttk.Style()
        s.theme_use("clam")
        s.configure("TProgressbar", thickness=14,
                     troughcolor=BORDER, background=ACCENT,
                     lightcolor=ACCENT, darkcolor=ACCENT, bordercolor=BORDER)

    # ── 控件工厂 ──────────────────────────────────────────────────────────────

    def _lbl(self, parent, text, size=9, bold=False, fg="#444", bg=None):
        return tk.Label(parent, text=text,
                        font=(FONT, size, "bold" if bold else ""),
                        bg=bg if bg is not None else parent["bg"], fg=fg)

    def _btn(self, parent, text, cmd, accent=False, green=False, **kw):
        bg = SUCCESS if green else (ACCENT if accent else "#e2e6ea")
        fg = "white" if (accent or green) else "#333"
        return tk.Button(parent, text=text, command=cmd,
                         bg=bg, fg=fg, activebackground=bg, activeforeground=fg,
                         relief="flat", bd=0, font=(FONT, 9),
                         padx=12, pady=5, cursor="hand2", **kw)

    def _card(self, parent, title):
        outer = tk.Frame(parent, bg=BORDER)
        inner = tk.Frame(outer, bg=CARD)
        inner.pack(fill="both", expand=True, padx=1, pady=1)
        hdr = tk.Frame(inner, bg="#f7f8fa")
        hdr.pack(fill="x")
        tk.Label(hdr, text=title, font=(FONT, 9, "bold"),
                 bg="#f7f8fa", fg="#555", anchor="w", padx=10, pady=7).pack(fill="x")
        tk.Frame(inner, height=1, bg=BORDER).pack(fill="x")
        body = tk.Frame(inner, bg=CARD, padx=12, pady=10)
        body.pack(fill="both", expand=True)
        return outer, body

    # ── 构建 UI ───────────────────────────────────────────────────────────────

    def _build_ui(self):
        wrap = tk.Frame(self.root, bg=BG)
        # 按钮固定在窗口最底部，不随内容多少移动
        self.go_btn = tk.Button(
            self.root, text="开始转换",
            command=self.start,
            bg=SUCCESS, fg="white",
            activebackground="#218838", activeforeground="white",
            relief="flat", bd=0,
            font=(FONT, 13, "bold"),
            pady=14,
            cursor="hand2",
        )
        self.go_btn.pack(side="bottom", fill="x")
        self._lbl(self.root, "LibreOffice / PyMuPDF / pdf2docx 三引擎",
                  size=7, fg="#aaa", bg=BG).pack(side="bottom", pady=(4, 0))

        wrap.pack(fill="both", expand=True, padx=18, pady=(14, 4))

        # 标题
        bar = tk.Frame(wrap, bg=BG)
        bar.pack(fill="x", pady=(0, 12))
        self._lbl(bar, "PDF 批量转 Word", size=15, bold=True, fg="#222", bg=BG).pack(side="left")
        dnd = "（支持拖拽）" if self._dnd_ok else ""
        self._lbl(bar, dnd, size=8, fg="#aaa", bg=BG).pack(side="left", padx=6)

        # ── 卡片 1：文件 ──────────────────────────────────────────────────────
        c1, b1 = self._card(wrap, "  选择 PDF 文件")
        c1.pack(fill="both", expand=True, pady=(0, 10))

        brow = tk.Frame(b1, bg=CARD)
        brow.pack(fill="x", pady=(0, 8))
        self._btn(brow, "+ 添加文件",   self.add_files,  accent=True).pack(side="left", padx=(0, 8))
        self._btn(brow, "+ 添加文件夹", self.add_folder, accent=True).pack(side="left", padx=(0, 8))
        self._btn(brow, "清空",         self.clear_list              ).pack(side="left")

        if self._dnd_ok:
            tk.Label(b1, text="可将 PDF 文件 / 文件夹拖拽到下方列表",
                     font=(FONT, 8), bg="#eef3fb", fg="#7a9ecf",
                     pady=5, anchor="center").pack(fill="x", pady=(0, 6))

        lb_wrap = tk.Frame(b1, bg=CARD)
        lb_wrap.pack(fill="both", expand=True)
        sy = tk.Scrollbar(lb_wrap); sy.pack(side="right", fill="y")
        sx = tk.Scrollbar(lb_wrap, orient="horizontal"); sx.pack(side="bottom", fill="x")

        self.lb = tk.Listbox(lb_wrap, font=(FONT, 9), height=7,
                              bg="#fafbfc", fg="#333",
                              selectbackground=ACCENT, selectforeground="white",
                              relief="flat", bd=1,
                              highlightthickness=1, highlightbackground=BORDER,
                              selectmode="extended", activestyle="none",
                              yscrollcommand=sy.set, xscrollcommand=sx.set)
        self.lb.pack(side="left", fill="both", expand=True)
        sy.config(command=self.lb.yview)
        sx.config(command=self.lb.xview)

        if self._dnd_ok:
            self.lb.drop_target_register(DND_FILES)
            self.lb.dnd_bind("<<Drop>>", self._on_drop)

        self.cnt_lbl = self._lbl(b1, "已选 0 个文件", size=8, fg="#aaa", bg=CARD)
        self.cnt_lbl.pack(anchor="w", pady=(5, 0))

        # ── 卡片 2：设置 ──────────────────────────────────────────────────────
        c2, b2 = self._card(wrap, "  输出设置")
        c2.pack(fill="x", pady=(0, 10))

        # 输出目录行
        r1 = tk.Frame(b2, bg=CARD); r1.pack(fill="x", pady=(0, 10))
        self._lbl(r1, "输出目录：", bg=CARD).pack(side="left")
        self.out_entry = tk.Entry(r1, textvariable=self.out_dir, font=(FONT, 9),
                                   relief="flat", bd=0,
                                   highlightthickness=1, highlightbackground="#d0d5de",
                                   bg=CARD, fg="#222", insertbackground="#222")
        self.out_entry.pack(side="left", fill="x", expand=True, padx=(4, 8))
        self._btn(r1, "浏览…", self.browse_out).pack(side="left")
        self._lbl(r1, "留空则与源文件同目录", size=8, fg="#bbb", bg=CARD).pack(side="left", padx=(8, 0))

        # 转换模式行（两个大按钮式单选）
        r2 = tk.Frame(b2, bg=CARD); r2.pack(fill="x")
        self._lbl(r2, "转换模式：", bg=CARD).pack(side="left")

        def _radio(parent, text, sub, val):
            f = tk.Frame(parent, bg=CARD, padx=2)
            f.pack(side="left", padx=(8, 0))
            tk.Radiobutton(f, text=text, variable=self.engine, value=val,
                           font=(FONT, 9, "bold"), bg=CARD, fg="#333",
                           activebackground=CARD, selectcolor=CARD,
                           cursor="hand2").pack(anchor="w")
            tk.Label(f, text=sub, font=(FONT, 7), bg=CARD, fg="#999",
                     wraplength=200, justify="left").pack(anchor="w")

        _radio(r2, "LibreOffice（最佳）",
               "表格/段落识别最准，需已安装 LibreOffice",
               "libreoffice")
        _radio(r2, "智能提取",
               "段落结构清晰，本地无需额外安装",
               "smart")
        _radio(r2, "高保真复刻",
               "尽量还原视觉排版",
               "fidelity")

        # ── 卡片 3：进度 ──────────────────────────────────────────────────────
        c3, b3 = self._card(wrap, "  转换进度")
        c3.pack(fill="x", pady=(0, 14))

        self.status = self._lbl(b3, "就绪，等待开始…", fg="#666", bg=CARD)
        self.status.pack(fill="x", anchor="w", pady=(0, 6))
        self.pbar = ttk.Progressbar(b3, mode="determinate")
        self.pbar.pack(fill="x")
        self.prog = self._lbl(b3, "0 / 0", size=8, fg="#aaa", bg=CARD)
        self.prog.pack(anchor="e", pady=(3, 0))

        # 占位弹性空间，让进度卡片不会过度拉伸
        tk.Frame(wrap, bg=BG).pack(fill="both", expand=True)

    # ── 文件管理 ──────────────────────────────────────────────────────────────

    def add_files(self):
        paths = filedialog.askopenfilenames(
            title="选择 PDF 文件",
            filetypes=[("PDF 文件", "*.pdf"), ("所有文件", "*.*")])
        if paths:
            self._push(list(paths))

    def add_folder(self):
        d = filedialog.askdirectory(title="选择包含 PDF 的文件夹")
        if not d:
            return
        found = sorted(Path(d).rglob("*.pdf"))
        if found:
            self._push([str(p) for p in found])
        else:
            messagebox.showinfo("提示", "该文件夹中没有找到 PDF 文件。")

    def browse_out(self):
        d = filedialog.askdirectory(title="选择输出目录")
        if d:
            self.out_dir.set(d)

    def clear_list(self):
        self.files.clear()
        self.lb.delete(0, "end")
        self._refresh_count()

    def _push(self, paths):
        for p in paths:
            if p not in self.files:
                self.files.append(p)
                self.lb.insert("end", p)
        self.lb.see("end")
        self._refresh_count()

    def _refresh_count(self):
        n = len(self.files)
        self.cnt_lbl.config(text=f"已选 {n} 个文件")

    def _on_drop(self, event):
        tokens = re.findall(r"\{[^}]+\}|[^\s]+", event.data)
        paths = []
        for t in tokens:
            t = t.strip("{}")
            p = Path(t)
            if p.is_dir():
                paths.extend(str(f) for f in sorted(p.rglob("*.pdf")))
            elif p.suffix.lower() == ".pdf" and p.is_file():
                paths.append(str(p))
        if paths:
            self._push(paths)
        else:
            messagebox.showinfo("提示", "拖入内容中未找到 PDF 文件。")

    # ── 转换 ──────────────────────────────────────────────────────────────────

    def start(self):
        if self.running:
            return
        if not self.files:
            messagebox.showwarning("提示", "请先添加 PDF 文件。")
            return
        self.running = True
        self.go_btn.config(state="disabled", text="转换中…", bg=GRAY)
        # 在主线程读取 tkinter 变量，再传给子线程（tkinter 非线程安全）
        engine  = self.engine.get()
        out_dir = self.out_dir.get().strip()
        threading.Thread(target=self._worker, args=(engine, out_dir), daemon=True).start()

    def _worker(self, engine: str, out_dir: str):
        files  = list(self.files)
        total  = len(files)
        ok, fail = 0, []

        self._after(lambda: self.pbar.config(maximum=total, value=0))

        for i, src in enumerate(files):
            name = Path(src).name
            self._update_progress(i, total, f"正在转换（{i+1}/{total}）：{name}")

            try:
                dest_dir = (Path(out_dir) if out_dir else Path(src).parent)
                dest_dir.mkdir(parents=True, exist_ok=True)
                dest = str(dest_dir / (Path(src).stem + ".docx"))

                if engine == "libreoffice":
                    convert_libreoffice(src, dest)
                elif engine == "smart":
                    convert_smart(src, dest)
                else:
                    from pdf2docx import Converter
                    cv = Converter(src)
                    cv.convert(dest, start=0, end=None)
                    cv.close()

                ok += 1
            except PermissionError:
                fail.append((name, "文件加密或拒绝访问"))
            except Exception as exc:
                fail.append((name, str(exc)))

        self._update_progress(total, total,
                               f"完成！成功 {ok} 个，失败 {len(fail)} 个。")
        self.running = False

        open_dir = out_dir or (str(Path(files[0]).parent) if files else "")
        self._after(lambda: self._show_done(ok, fail, open_dir))

    def _update_progress(self, cur, total, text):
        def _do():
            self.status.config(text=text)
            self.pbar.config(value=cur)
            self.prog.config(text=f"{cur} / {total}")
        self._after(_do)

    def _after(self, fn):
        self.root.after(0, fn)

    def _show_done(self, ok, fail, open_dir):
        self.go_btn.config(state="normal", text="开始转换", bg=SUCCESS)
        lines = [f"转换完成！\n\n成功：{ok} 个\n失败：{len(fail)} 个"]
        if fail:
            lines.append("\n\n失败文件详情：")
            for name, err in fail:
                short = (err[:120] + "…") if len(err) > 120 else err
                lines.append(f"\n• {name}\n  {short}")
        msg = "".join(lines)
        if open_dir:
            if messagebox.askquestion("转换完成", msg + "\n\n是否打开输出文件夹？",
                                       icon="info") == "yes":
                self._open_dir(open_dir)
        else:
            messagebox.showinfo("转换完成", msg)

    @staticmethod
    def _open_dir(path: str):
        s = platform.system()
        if s == "Windows":   os.startfile(path)
        elif s == "Darwin":  subprocess.Popen(["open", path])
        else:                subprocess.Popen(["xdg-open", path])


if __name__ == "__main__":
    App()
